"""
Document generator for employee documents (Employment Agreement, etc.).
Builds context from employee + company + template and renders body with Django Template.
"""

from datetime import date
from io import BytesIO

from django.template import Context, Template

from base.models import Company, DocumentTemplate
from employee.models import Employee


def get_document_context(employee, document_template, agreement_date=None):
    """
    Build context dict for rendering a document template with employee and company data.
    """
    if agreement_date is None:
        agreement_date = date.today()

    work_info = getattr(employee, "employee_work_info", None)
    company = None
    if work_info and getattr(work_info, "company_id", None):
        company = work_info.company_id
    if not company and document_template.company_id:
        company = document_template.company_id
    if not company:
        company = Company.objects.filter(hq=True).first()

    # Company fields
    company_name = company.company if company else ""
    company_address = ""
    if company:
        parts = [company.address or "", company.city or "", company.state or "", company.zip or "", company.country or ""]
        company_address = ", ".join(p for p in parts if p).strip(", ")

    # Employee fields
    job_position = ""
    if work_info and work_info.job_position_id:
        job_position = str(work_info.job_position_id.job_position)
    joining = work_info.date_joining if work_info else None
    employment_commencement_date = joining.strftime("%d/%m/%Y") if joining else agreement_date.strftime("%d/%m/%Y")

    # Employee address: address, city, state, zip, country
    addr_parts = [employee.address or "", employee.city or "", employee.state or "", employee.zip or "", employee.country or ""]
    employee_address = ", ".join(p for p in addr_parts if p).strip(", ")

    # Work location: use company address as office address
    work_location = company_address or (getattr(work_info, "location", None) or "")

    # Signatory from template or empty
    signatory_name = document_template.signatory_name or ""
    signatory_designation = document_template.signatory_designation or ""

    context = {
        "agreement_date": agreement_date.strftime("%d/%m/%Y"),
        "company_name": company_name,
        "company_address": company_address,
        "employee_name": employee.get_full_name(),
        "employee_first_name": employee.employee_first_name,
        "employee_last_name": employee.employee_last_name or "",
        "employee_address": employee_address,
        "employee_email": employee.email or "",
        "employee_phone": employee.phone or "",
        "job_position": job_position,
        "employment_commencement_date": employment_commencement_date,
        "work_location": work_location,
        "probation_months": "3",
        "signatory_name": signatory_name,
        "signatory_designation": signatory_designation,
    }
    return context


def render_document_body(document_template, context):
    """Render template body with context; return plain text (no HTML markup)."""
    body = document_template.body or ""
    t = Template(body)
    ctx = Context(context)
    return t.render(ctx)


def _replace_placeholders_in_text(text, context):
    """
    Replace {{ key }} or {{key}} in text with context values.
    Supports spaces inside braces and normalizes key (strip, spaces -> underscores) for lookup.
    """
    import re
    if not text or "{{" not in text:
        return text
    # Build lookup: exact key, and with spaces instead of underscores (e.g. "employee name" -> employee_name's value)
    context_lookup = {k: str(v) for k, v in context.items()}
    for k in list(context.keys()):
        context_lookup[k.replace("_", " ")] = str(context[k])
    # Match {{ ... }} with optional spaces around the key
    pattern = re.compile(r"\{\{\s*([^}]+)\s*\}\}")
    def repl(match):
        key_raw = match.group(1).strip()
        key_underscore = key_raw.replace(" ", "_")
        if key_raw in context_lookup:
            return context_lookup[key_raw]
        if key_underscore in context_lookup:
            return context_lookup[key_underscore]
        if key_raw in context:
            return str(context[key_raw])
        return match.group(0)
    return pattern.sub(repl, text)


def _get_paragraph_text(paragraph):
    """Get full text of paragraph (in case Word split content across runs)."""
    return "".join(run.text for run in paragraph.runs)


def _get_paragraph_font(paragraph):
    """Read font properties from first run (before any clear). Returns dict to apply later."""
    saved = {}
    if not paragraph.runs:
        return saved
    r = paragraph.runs[0]
    try:
        if r.font.name:
            saved["name"] = r.font.name
        if r.font.size:
            saved["size"] = r.font.size
        saved["bold"] = r.font.bold
        saved["italic"] = r.font.italic
        saved["underline"] = r.font.underline
    except Exception:
        pass
    return saved


def _apply_font_to_run(run, font_saved):
    """Apply saved font properties to a run."""
    if not font_saved:
        return
    try:
        if font_saved.get("name"):
            run.font.name = font_saved["name"]
        if font_saved.get("size"):
            run.font.size = font_saved["size"]
        if "bold" in font_saved:
            run.font.bold = font_saved["bold"]
        if "italic" in font_saved:
            run.font.italic = font_saved["italic"]
        if "underline" in font_saved:
            run.font.underline = font_saved["underline"]
    except Exception:
        pass


def _replace_in_paragraph(paragraph, context, replace_count_list):
    """Replace placeholders in one paragraph. Preserves first run's font. Tracks if any replacement was made."""
    text = _get_paragraph_text(paragraph)
    if not text:
        return
    new_text = _replace_placeholders_in_text(text, context)
    if new_text != text:
        font_saved = _get_paragraph_font(paragraph)
        paragraph.clear()
        new_run = paragraph.add_run(new_text)
        _apply_font_to_run(new_run, font_saved)
        replace_count_list[0] += 1


def _replace_literal_fallback_in_paragraph(paragraph, context, literal_map):
    """When template has no {{ }} placeholders, replace known sample text. Preserves font."""
    text = _get_paragraph_text(paragraph)
    if not text:
        return
    new_text = text
    for sample_literal, context_key in literal_map.items():
        if sample_literal in new_text and context_key in context:
            value = context[context_key]
            if value is None or str(value).strip() == "":
                value = "—"
            new_text = new_text.replace(sample_literal, str(value))
    if new_text != text:
        font_saved = _get_paragraph_font(paragraph)
        paragraph.clear()
        new_run = paragraph.add_run(new_text)
        _apply_font_to_run(new_run, font_saved)


def _all_paragraphs(doc):
    """Yield all paragraphs in document body, tables, headers, footers."""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for cell in table.cells:
            for para in cell.paragraphs:
                yield para
    for section in doc.sections:
        for para in section.header.paragraphs:
            yield para
        for para in section.footer.paragraphs:
            yield para


def generate_document_from_uploaded_docx(document_template, employee, agreement_date=None):
    """
    Generate document from an uploaded Word (.docx) template file.
    Replaces {{ placeholder }} in all paragraphs, table cells, headers and footers.
    If the file has no placeholders, replaces common sample text (e.g. Darshan P., Vehicle Driver) with selected employee data.
    """
    if not document_template.template_file:
        raise ValueError("This template has no uploaded file. Use text template or upload a .docx file.")

    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    from django.core.files.storage import default_storage

    context = get_document_context(employee, document_template, agreement_date)

    path = document_template.template_file.name
    with default_storage.open(path, "rb") as f:
        doc = Document(f)

    replace_count = [0]  # mutable so we can update inside nested function

    # First pass: replace {{ placeholder }} in every paragraph
    for para in _all_paragraphs(doc):
        _replace_in_paragraph(para, context, replace_count)

    # Fallback: if no placeholders were found, replace common sample literals (template was pre-filled with sample data)
    if replace_count[0] == 0:
        literal_map = {
            "Darshan P.": "employee_name",
            "Darshan P": "employee_name",
            "Vehicle Driver": "job_position",
            "02/03/2026": "agreement_date",
            # Sample addresses (multiple variants in case of hyphen or spacing)
            "S/O: Prakasha, Thattekere Village, Kanakapura Taluk, Ramanagar, Karnataka - 562112": "employee_address",
            "S/O: Prakasha, Thattekere Village, Kanakapura Taluk, Ramanagar, Karnataka 562112": "employee_address",
            "Thattekere Village, Kanakapura Taluk, Ramanagar, Karnataka - 562112": "employee_address",
            "Thattekere Village, Kanakapura Taluk, Ramanagar, Karnataka 562112": "employee_address",
        }
        for para in _all_paragraphs(doc):
            _replace_literal_fallback_in_paragraph(para, context, literal_map)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _build_agreement_html(body_text):
    """
    Build structured HTML from agreement text so the PDF matches a formal
    contract format: title, date, parties, numbered sections and sub-clauses.
    """
    import html
    import re

    lines = body_text.split("\n")
    out = []
    # Section header: "1. TITLE" or "2. PROBATION"
    section_re = re.compile(r"^(\d+)\.\s+(.+)$")
    # Sub-clause: "1.1 Text" or "2.3 Text"
    clause_re = re.compile(r"^(\d+)\.(\d+)\s+(.+)$")

    i = 0
    while i < len(lines):
        line = lines[i]
        escaped = html.escape(line.strip())
        if not escaped:
            i += 1
            continue
        # First non-empty line as document title
        if not out and len(line.strip()) < 80 and (line.strip().isupper() or "AGREEMENT" in line.strip().upper()):
            out.append(f'<h1 class="doc-title">{escaped}</h1>')
            i += 1
            continue
        # Sub-clause first (e.g. "1.1 The Company appoints...") so we don't treat as section
        m = clause_re.match(line.strip())
        if m:
            out.append(f'<p class="clause">{escaped}</p>')
            i += 1
            continue
        # Numbered section header (e.g. "1. APPOINTMENT & COMMENCEMENT")
        m = section_re.match(line.strip())
        if m:
            rest = m.group(2)
            if not rest or len(rest) < 120:
                out.append(f'<p class="section-header">{escaped}</p>')
                i += 1
                continue
        # BETWEEN / AND labels
        if line.strip().upper() in ("BETWEEN", "AND"):
            out.append(f'<p class="party-label">{escaped}</p>')
            i += 1
            continue
        # Bullet lines (•) or normal paragraphs
        if line.strip().startswith("•") or line.strip().startswith("-"):
            out.append(f'<p class="bullet">{escaped}</p>')
        else:
            out.append(f'<p class="body">{escaped}</p>')
        i += 1

    return "\n".join(out) if out else html.escape(body_text).replace("\n", "<br/>")


def generate_document_pdf(document_template, employee, agreement_date=None):
    """
    Render the document template for the employee and return PDF bytes.
    Uses xhtml2pdf with structured HTML so the output matches a formal
    Employment Agreement format (title, sections, sub-clauses).
    """
    from xhtml2pdf import pisa

    context = get_document_context(employee, document_template, agreement_date)
    body_text = render_document_body(document_template, context)
    body_html = _build_agreement_html(body_text)

    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8"/>
        <style>
            body {{
                font-family: Arial, sans-serif;
                font-size: 11pt;
                line-height: 1.45;
                margin: 24px 28px;
                color: #222;
            }}
            .doc-title {{
                font-size: 16pt;
                font-weight: bold;
                text-align: center;
                margin-bottom: 14pt;
            }}
            .section-header {{
                font-weight: bold;
                font-size: 11pt;
                margin-top: 10pt;
                margin-bottom: 4pt;
            }}
            .clause {{
                margin: 2pt 0 2pt 0;
                text-align: justify;
            }}
            .party-label {{
                font-weight: bold;
                margin-top: 8pt;
                margin-bottom: 2pt;
            }}
            .bullet {{
                margin: 1pt 0 1pt 12pt;
            }}
            .body {{
                margin: 2pt 0;
            }}
        </style>
    </head>
    <body>
    {body_html}
    </body>
    </html>
    """
    result = BytesIO()
    status = pisa.CreatePDF(src=html_content, dest=result, encoding="utf-8")
    if status.err:
        raise ValueError("PDF generation failed")
    return result.getvalue()


def generate_document_docx(document_template, employee, agreement_date=None):
    """
    Render the document template for the employee and return DOCX bytes.
    Requires python-docx.
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise ImportError("python-docx is required for Word export. Install with: pip install python-docx")

    context = get_document_context(employee, document_template, agreement_date)
    body_text = render_document_body(document_template, context)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Arial"

    # Split by double newline for paragraphs, then by single newline for lines
    blocks = body_text.split("\n\n")
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = block.split("\n")
        for i, line in enumerate(lines):
            p = doc.add_paragraph(line.strip())
            if i < len(lines) - 1:
                p.add_run().add_break()
    if not list(doc.paragraphs):
        doc.add_paragraph(body_text)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
